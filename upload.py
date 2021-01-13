#OS
import os
#Doc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
#Pandas
import pandas as pd,numpy as np,sqlite3
#Flask API
from flask import Flask, render_template,request,send_file,redirect,url_for,after_this_request,send_from_directory,abort,flash
from flask_cors import CORS
from werkzeug.utils import secure_filename
import json,ast
from functools import reduce
from funcs import *
from io import StringIO
import time
from datetime import datetime
import time
from docx.shared import RGBColor,Pt,Inches
from docx_utils.flatten import opc_to_flat_opc
#from docx2pdf import convert
import docx
from waitress import serve

def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


    

#Only csv files are allowed
ALLOWED_DOCS = {'csv','xlsx'}
ALLOWED_IMAGES = {'png','jpg','jpeg'}
File_Name = {'Template'}
app = Flask(__name__)
app.secret_key = 'random string'

#Give always a new file for export (as we are using send_file)
app.config['SEND_FILE_MAX_AGE_DEFAULT']=0 
#Size 16MB
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
#Cache control
CORS(app) 
app.config['TEMPLATES_AUTO_RELOAD'] = True



@app.after_request
def add_header(response):
    # response.cache_control.no_store = True
    if 'Cache-Control' not in response.headers:
        response.headers['Cache-Control'] = 'no-store'
   
    return response

#Current Directory


	
@app.route('/<project>/delete', methods = ['GET','POST'])
def delete_files(project):
   if request.method=='GET':
    try:
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.db')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'.db'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.json')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'.json'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'_content.csv'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'_sort.json')) 
     flash('Project has been deleted')
     return redirect('/')
    except Exception as e:

     flash("Project couldn't be deleted")
     return redirect('/')


	
@app.route('/<project>', methods = ['GET','POST'])
def upload_files(project):
   if request.method=='GET':  
    try:
       if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
        with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
         json_sort = json.load(json_file)
        sheets_needed=list(json_sort.keys()  )
   
       else:
          sheets_needed=None  
          
     
       input_file  = db_read(project,'input',sheets_needed=sheets_needed)
    
       colummnsList = input_file
     
       
              
     
       colummnsList = ['{'+ x +'}' for x in colummnsList]

       df = db_read(project,'template')
   

       df=df[["Sno","Business_Rule","String_Name"]]
       df=df.replace(np.nan,'',regex=True)
       items = df.to_dict('records')

      
             
       if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
 
       
         json_decoded={'All':''} 
         
         with open(os.path.join(os.getcwd(),'Data',project+'_sort.json'),'w') as json_file:
              json.dump(json_decoded, json_file)              


       if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.json')):
 
        with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
         json_decoded = json.load(json_file)
     
         
       else:
         json_decoded={'key':'#sp','header':'','description':'','datadefinition':'','datatemplate':'','rowbinder':'','format':'#word'} 
         
         with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
              json.dump(json_decoded, json_file)
       return render_template("Lat.html",items=items,columns=colummnsList,
                              project=project,thing=json_decoded['key'],
                              datadefinition=json_decoded['datadefinition'],
                              datatemplate=json_decoded['datatemplate'],
                              rbind=json_decoded['rowbinder'],
                              fmat=json_decoded['format'])

    except Exception as e:
 
      
  
      return render_template("Lat.html",project=project)  
    
 
   if request.method == 'POST':
    try:
     
      req = request.get_json()

      df=pd.DataFrame(req)
      df.columns = ["Sno","Business_Rule","String_Name"]

      db_store(project,dataframe=df,uploads='template')
      controller(project)
   
      return redirect(f'/{project}')
    except Exception as e:
           
           return redirect(f'/{project}')


@app.route('/<project>/configure', methods = ['POST'])
def add_config(project):

    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file) 

    except Exception as e:

     json_decoded={} 

    json_decoded['key'] = request.form['Paragraph_Config']
  


    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)
    
    controller(project)
    return redirect(f'/{project}') 

@app.route('/<project>/rowbinder', methods = ['POST'])
def add_rowbinder(project):
    
    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file) 
    except Exception as e:
     
     json_decoded={} 

    json_decoded['rowbinder'] = request.form['rowbinder']
  


    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)
    
    controller(project)
    return redirect(f'/{project}') 

@app.route('/<project>/format', methods = ['POST'])
def add_format(project):
    
    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file) 
    except Exception as e:
     
     json_decoded={} 

    json_decoded['format'] = request.form['format']
  


    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)
    
    #controller(project)
    return redirect(f'/{project}') 

@app.route('/<project>/sortby', methods = ['GET','POST'])
def sortby(project):
    
    if request.method == 'GET':
        
        dbname = project
        conn = sqlite3.connect(os.path.join(os.getcwd(),"Data",dbname+'.db'))
        f=[','.join(x) for x in conn.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()]
        f=[x for x in f if x!='template']
        b = {}
   
        if f!=[]:
         for db in f:  
          df=pd.read_sql(f'select * from {db}',conn).drop('index',axis=1)

          b[db] = df.columns.tolist()
      
         
         sheet = [(k, v) for k, v in b.items()]
         if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
             with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
                 proj_json = json.load(json_file)
                 
         return render_template("Sortby.html",project=project,sheets=sheet,sheets_selected=list(proj_json.keys()),columns_selected=proj_json.values())
        else:
         return render_template("Sortby.html",project=project,sheets=[],sheets_selected=[],columns_selected=[])
         
     
     
    if request.method == 'POST':
 
        
  
     
        d=dict([ [(x,a) for a in request.form.getlist(x) if a!='on'][0] for x in [x[0] for x in [tup for tup in request.form.items() if any(i in tup for i in ['on'])]]])
        
        with open(os.path.join(os.getcwd(),'Data',project+'_sort.json'),'w') as json_file:
            json.dump(d, json_file)
        
        
 
 
      
        dbname = project
        conn = sqlite3.connect(os.path.join(os.getcwd(),"Data",dbname+'.db'))
        f=[','.join(x) for x in conn.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()]
        f=[x for x in f if x!='template']
        b = {}
        if f!=[]:
         for db in f:  
          df=pd.read_sql(f'select * from {db}',conn).drop('index',axis=1)
          #b["SheetName"] = db
          b[db] = df.columns.tolist()
          #b["ColumnList"] = df.columns.tolist()
          #a = {"SheetName":db,"ColumnList":df.columns.tolist()}
          #b = json.loads(a).update(a)/

         
         sheet = [(k, v) for k, v in b.items()]
        controller(project) 
       
        with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
                 proj_json = json.load(json_file)
          
        return render_template("Sortby.html",project=project,sheets=sheet,sheets_selected=list(proj_json.keys()),columns_selected=list(proj_json.values()))
        

    
@app.route('/<project>/<uploads>', methods = ['POST'])
def upload_input(project,uploads):


 
   if request.method == 'POST':
    try:
      #f = request.files['file'] 

      uploaded_files = request.files

      uploaded_files = uploaded_files.to_dict(flat=False)
 
      files = uploaded_files["file"]
      

      if set([file.filename for file in files])!={''}:
       for file in files:
          if (file.filename).split(".")[1] in ALLOWED_DOCS:
              file.save(os.path.join('temp',secure_filename(file.filename)))
     
         
              db_store(project,filename=file.filename,uploads=uploads)
              
              controller(project)
         
             
          else:
       
              return redirect(f'/{project}')
  

       flash('File has been uploaded successfully')
       return redirect(f'/{project}')
      else:  
       return redirect(f'/{project}')
    except Exception as e:
  
      return redirect(f'/{project}')
   

#Home page 
@app.route('/', methods = ['GET','POST'])
def upload_file():
   if request.method=='GET' :
       df=pd.DataFrame([x.split('.')[0] for x in os.listdir('Data') if 'favicon' not in x and '.db' in x])
  
       if not df.empty:
        df.columns=['columns']
        df['sno']=['val'+str(x) for x in np.arange(len(df))+1]
        df['links']=request.host_url+df['columns']
      
        df=df.replace(np.nan,'',regex=True).drop_duplicates()
        items = df.to_dict('records')
   
       else:
         items={}
       if os.path.exists(os.path.join('Data','logo.png')):
          logo='yes'
       else:
         logo=''       
       return render_template('Multi.html',items=items,logo=logo)   
       
   if request.method=='POST':
  
    jsonresponse=request.form.to_dict()

  
    if jsonresponse=={}: 
  
       flash('Nothing has been selected')
       return redirect('/')
    elif 'textbox' not  in jsonresponse.keys() :
     try:
      l=[x[1] for x in list(jsonresponse.items()) if 'chk[]' not in x[0]]
 
      for i in l:
       controller(i)
      document = Document()
      if os.path.exists(os.path.join(os.getcwd(),'Data','logo.png')):
       document.add_picture(os.path.join(os.getcwd(),'Data','logo.png'))
       last_paragraph = document.paragraphs[-1] 
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
      
      #Project json
      with open(os.path.join(os.getcwd(),'Data','Project.json')) as json_file:
         json_decoded_consolidated = json.load(json_file) 
      #run = document.add_paragraph().add_run()
      
      style = document.styles['Normal']
      font = style.font
      font.name = 'Calibri'
      font.size = Pt(28)
      font.color.rgb = RGBColor(0,0,0)
      font.bold = True
      paragraph1 = document.add_paragraph(json_decoded_consolidated['Projectheader']+"\n")
      #run = paragraph1.add_run()
      #font = run.font
      #font.name = 'Calibri'
      #font.size = Pt(20)
      #font.color.rgb = RGBColor(0,0,0)
      #font.bold = True
      paragraph_format = paragraph1.paragraph_format
      paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
      
      sentence = paragraph1.add_run(json_decoded_consolidated['Projectdescription']+"\n")  
      sentence.font.size = Pt(20)
      sentence.font.color.rgb = RGBColor(192,192,192)
      

      
      for i in l:
       df = pd.read_csv(os.path.join(os.getcwd(),'Data',i+'_content.csv'))
      
       with open(os.path.join(os.getcwd(),'Data',i+'.json')) as json_file:
          json_decoded = json.load(json_file) 
       #run = document.add_paragraph().add_run()
        
       style = document.styles['Normal']
       font = style.font
       font.name = 'Calibri'
       font.size = Pt(13)
       font.color.rgb = RGBColor(0,105,225)
       font.bold = True
       paragraph=document.add_paragraph(json_decoded['header']+"\n")
       sentence = paragraph.add_run(json_decoded['description']+"\n")  
       sentence.font.size=Pt(11)
       sentence.font.italic = True
       for index,row in df.iterrows():
      
         sentence = paragraph.add_run(row.vals+"\n")
         sentence.font.name = 'Calibri'
         sentence.font.bold = False
         sentence.font.size = Pt(11)
         sentence.font.color.rgb=RGBColor(0,0,0)
         
      
      
      
      
      document.save(os.path.join(os.getcwd(),'temp','Combo.docx'))
   
      return send_from_directory('temp','Combo.docx',as_attachment=True,attachment_filename=f'Combo_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.docx')
     except Exception as e:
      flash('File could not be downloaded')
      return redirect('/')
    else:
     if jsonresponse['textbox'].lower() not in [x.split('.')[0].lower() for x in os.listdir('Data') if 'favicon' not in x]:
    
      return redirect(f"{request.host_url}{jsonresponse['textbox']}")
     else:
       flash('Project already exists with the same name')
       return redirect('/')

#download project level output here  
@app.route('/<project>/output',methods = ['GET'])
def download_file(project):
    if request.method == 'GET':
     try: 
       
       
       
       if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):
       
        controller(project)
    
       df=pd.read_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))
    
       document=Document()
       
       
       with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
         json_decoded = json.load(json_file) 
       #run = document.add_paragraph().add_run()
       
       style = document.styles['Normal']
       font = style.font
       font.name = 'Calibri'
       
       font.size = Pt(13)
       #font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
       font.color.rgb = RGBColor(0,105,225)
       font.bold = True
       
       #download_format = json_decoded['format']
       paragraph = document.add_paragraph(json_decoded['header']+'\n')
       sentence = paragraph.add_run(json_decoded['description']+"\n")  
       sentence.font.size=Pt(11)
       sentence.font.italic = True
       for index,row in df.iterrows():
     
        sentence = paragraph.add_run(row.vals+"\n")
        sentence.font.bold = False
        sentence.font.name = 'Calibri'
        sentence.font.size = Pt(11)
        sentence.font.color.rgb=RGBColor(0,0,0)
         
       
       document.save(os.path.join(os.getcwd(),'temp',f'{project}.docx'))
       #convert(os.path.join(os.getcwd(),'temp',f'{project}.docx'),os.path.join(os.getcwd(),'temp',f'{project}.pdf'))
       #if download_format == '#word':
       return send_from_directory('temp',f'{project}.docx',as_attachment=True,attachment_filename=f'{project}_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.docx')
       #elif download_format == '#pdf':
       #return send_from_directory('temp',f'{project}.pdf',as_attachment=True,attachment_filename=f'{project}_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.pdf')
       
     except Exception as e:
  
       return redirect(f'/{project}')
   
@app.route('/<project>/templatedownload',methods = ['GET'])
def templatedown_file(project):
    if request.method == 'GET':
     try:
      
       df=db_read(project,'template')

       df.to_csv(os.path.join(os.getcwd(),'temp',f'{project}_template.csv'),index=False) 
 
       return send_from_directory('temp',f'{project}_template.csv',as_attachment=True,attachment_filename=f'{project}_template_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.csv')
     
     except Exception as e:
      
       return redirect(f'/{project}')

@app.route('/<project>/description', methods = ['POST'])
def add_header(project):
    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file) 
    except:
     json_decoded={} 

    json_decoded['header'] = request.form['Header']
    json_decoded['description']=request.form['Description']


    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)
 
    return redirect(f'/{project}')

@app.route('/projectdescription', methods = ['POST'])
def add_project_header():
    try:
     with open(os.path.join(os.getcwd(),'Data','project.json')) as json_file:
      json_decoded = json.load(json_file) 
    except:
     json_decoded={} 

    json_decoded['Projectheader'] = request.form['ProjectHeader']
    json_decoded['Projectdescription']=request.form['ProjectDescription']


    with open(os.path.join(os.getcwd(),'Data','Project.json'),'w') as json_file:
     json.dump(json_decoded, json_file)
 
    return redirect(f'/')

@app.route('/logo', methods = ['POST'])
def upload_logo():
 
   if request.method == 'POST':
    try:
      #f = request.files['file'] 

      uploaded_files = request.files
  
      uploaded_files = uploaded_files.to_dict(flat=False)
     
      files = uploaded_files["file"]
      if set([file.filename for file in files])!={''}:
       for file in files:
          if (file.filename).split(".")[1].lower() in ALLOWED_IMAGES:
              file.save(os.path.join('Data',secure_filename('logo.png')))
     
      flash('The logo has been uploaded ')
      return redirect('/')
    except Exception as e:
   
     return redirect('/')
     

@app.route('/delete_logo',methods=['GET'])
def delete_logo():

 if os.path.exists(os.path.join('Data','logo.png')):
 
   os.remove(os.path.join('Data','logo.png'))
   flash('The logo has been deleted successfully')
   return redirect('/')
 else:
  flash('There is no logo present')
  return redirect('/')
  
   
     
     
@app.route('/<project>/static', methods = ['GET'])
def read_header(project):

      with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
          json_decoded = json.load(json_file)
      return json_decoded  

@app.route('/projectstatic', methods = ['GET'])
def read_project_header():

      with open(os.path.join(os.getcwd(),'Data','Project.json')) as json_file:
          json_decoded = json.load(json_file)
      return json_decoded       
    

    
@app.route('/<project>/preview', methods = ['GET','POST'])
def preview_file(project):
    if request.method=='GET':
      if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):
       
        controller(project)
      with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
          json_decoded = json.load(json_file)
      df=pd.read_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))
     
      return {"Header":json_decoded['header'],"Description":json_decoded['description'],"Content":'\n'.join(df['vals'].tolist() )} 	
    if request.method=='POST': 
       a = request.form.to_dict()
       with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
          json_decoded = json.load(json_file)
     
       json_decoded['header']=a['Header']
       json_decoded['description']=a['Description']
       with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
            json.dump(json_decoded, json_file)
       pd.DataFrame(a['Content'].split('\n'),columns=['vals']).to_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))
       
       return redirect(f'/{project}') 
         

######################################################################
############################ New Enhancements#########################
######################################################################

#uplpoad files here
@app.route('/Resolution/<uploads>', methods = ['POST'])
def upload_ResolutionFiles(uploads):
    if request.method == 'POST':
        
        try:
            
            uploaded_files = request.files.to_dict(flat=False)["file"]
            #uploaded_files = uploaded_files.to_dict(flat=False)
            #files = uploaded_files["file"]

            for file in uploaded_files:

                #if file.filename.split('.')[-1] == 'docx':
                    
                #    print('hey',file.filename)
                #elif file.filename.split('.')[-1] == 'xlsx':
                #    print('hi',file.filename)
                
                if uploads == 'word':
                    print('hi')
                    file.save(os.path.join('temp',secure_filename('word.docx')))
                    #doc = docx.Document(os.path.join(os.getcwd(),'sample.xml'))
                    print('2')
                    doc = docx.Document(os.path.join(os.getcwd(),'temp','word.docx'))
                    print('1')
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    print(len(fullText))
                    

            
            
            return redirect('/')
        
        except Exception as e:
            print(str(e))
            return redirect('/')
     
if __name__ == '__main__':
   serve(app,port=5000) 
